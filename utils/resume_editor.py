#!/usr/bin/env python3
"""
Resume Editor Utility

This module provides comprehensive tools for editing resumes to match job descriptions.
It includes AI-powered content optimization, ATS optimization, and file manipulation capabilities.
"""

import os
import re
import json
import asyncio
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import openai
from openai import AsyncOpenAI
from config import Config

class ResumeEditor:
    """Comprehensive resume editing utility with AI-powered optimization."""
    
    def __init__(self, api_key: Optional[str] = None):
        """Initialize the ResumeEditor."""
        self.api_key = api_key or Config.OPENAI_API_KEY
        self.client = AsyncOpenAI(api_key=self.api_key) if self.api_key else None
        
    async def edit_resume_for_job(self, 
                                resume_path: str,
                                job_title: str,
                                job_description: str,
                                company_name: str,
                                output_dir: Optional[str] = None) -> Dict[str, Any]:
        """
        Edit a resume to match a specific job description.
        
        Args:
            resume_path: Path to the original resume
            job_title: Target job title
            job_description: Job description text
            company_name: Company name
            output_dir: Output directory for modified resume
            
        Returns:
            Dictionary containing modification results and file path
        """
        
        try:
            # Validate inputs
            if not os.path.exists(resume_path):
                raise FileNotFoundError(f"Resume file not found: {resume_path}")
            
            if not job_title or not job_description:
                raise ValueError("Job title and description are required")
            
            # Extract resume content
            resume_content = self._extract_resume_content(resume_path)
            
            # Analyze job requirements
            job_analysis = self._analyze_job_requirements(job_description)
            
            # Analyze current resume
            resume_analysis = self._analyze_current_resume(resume_content)
            
            # Identify skill gaps
            skill_gaps = self._identify_skill_gaps(resume_analysis, job_analysis)
            
            # Generate optimized content
            optimized_content = await self._generate_optimized_content(
                resume_analysis, job_analysis, skill_gaps, job_title
            )
            
            # Create modification plan
            modification_plan = self._create_modification_plan(
                resume_analysis, job_analysis, skill_gaps, optimized_content
            )
            
            # Generate ATS recommendations
            ats_recommendations = self._generate_ats_recommendations(
                resume_analysis, job_analysis, skill_gaps
            )
            
            # Create modified resume file
            modified_resume_path = await self._create_modified_resume(
                resume_path, optimized_content, job_title, company_name, output_dir
            )
            
            return {
                "status": "success",
                "job_title": job_title,
                "company_name": company_name,
                "original_resume_path": resume_path,
                "modified_resume_path": modified_resume_path,
                "resume_analysis": resume_analysis,
                "job_analysis": job_analysis,
                "skill_gaps": skill_gaps,
                "optimized_content": optimized_content,
                "modification_plan": modification_plan,
                "ats_recommendations": ats_recommendations,
                "modification_timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "modification_timestamp": datetime.now().isoformat()
            }
    
    def _extract_resume_content(self, resume_path: str) -> str:
        """Extract text content from resume file."""
        
        try:
            if resume_path.endswith('.docx'):
                return self._extract_from_docx(resume_path)
            elif resume_path.endswith('.pdf'):
                return self._extract_from_pdf(resume_path)
            elif resume_path.endswith('.txt'):
                return self._extract_from_txt(resume_path)
            else:
                raise ValueError(f"Unsupported file format: {resume_path}")
        except Exception as e:
            raise Exception(f"Failed to extract content from {resume_path}: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        
        doc = Document(file_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content.append(cell.text.strip())
        
        return '\n'.join(content)
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = []
                for page in reader.pages:
                    content.append(page.extract_text())
                return '\n'.join(content)
        except ImportError:
            raise Exception("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _analyze_job_requirements(self, job_description: str) -> Dict[str, Any]:
        """Analyze job description to extract requirements."""
        
        # Extract technical skills
        technical_skills = self._extract_technical_skills(job_description)
        
        # Extract soft skills
        soft_skills = self._extract_soft_skills(job_description)
        
        # Extract experience requirements
        experience_requirements = self._extract_experience_requirements(job_description)
        
        # Extract education requirements
        education_requirements = self._extract_education_requirements(job_description)
        
        # Extract keywords
        keywords = self._extract_keywords(job_description)
        
        return {
            "technical_skills": technical_skills,
            "soft_skills": soft_skills,
            "experience_requirements": experience_requirements,
            "education_requirements": education_requirements,
            "keywords": keywords,
            "total_requirements": len(technical_skills) + len(soft_skills)
        }
    
    def _extract_technical_skills(self, text: str) -> List[Dict[str, Any]]:
        """Extract technical skills from text."""
        
        # Define skill categories and patterns
        skill_patterns = {
            "programming_languages": [
                r"python", r"java", r"javascript", r"typescript", r"c\+\+", r"c#", r"go", r"rust", r"php", r"ruby",
                r"swift", r"kotlin", r"scala", r"r", r"matlab", r"sql", r"html", r"css", r"bash", r"powershell"
            ],
            "frameworks": [
                r"django", r"flask", r"fastapi", r"react", r"angular", r"vue", r"node\.js", r"express", r"spring",
                r"laravel", r"rails", r"asp\.net", r"dotnet", r"flutter", r"react native", r"xamarin"
            ],
            "databases": [
                r"postgresql", r"mysql", r"mongodb", r"redis", r"elasticsearch", r"dynamodb", r"sqlite", r"oracle",
                r"sql server", r"mariadb", r"cassandra", r"neo4j", r"influxdb"
            ],
            "cloud_platforms": [
                r"aws", r"amazon web services", r"azure", r"gcp", r"google cloud", r"docker", r"kubernetes",
                r"terraform", r"ansible", r"chef", r"puppet", r"jenkins", r"gitlab", r"github actions"
            ],
            "tools": [
                r"git", r"jira", r"confluence", r"slack", r"figma", r"sketch", r"adobe", r"postman", r"swagger",
                r"kibana", r"grafana", r"prometheus", r"nagios", r"zabbix"
            ]
        }
        
        extracted_skills = []
        
        for category, patterns in skill_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, text.lower())
                for match in matches:
                    skill_name = match.group()
                    context_start = max(0, match.start() - 50)
                    context_end = min(len(text), match.end() + 50)
                    context = text[context_start:context_end].lower()
                    
                    # Calculate confidence based on context
                    confidence = 0.7
                    if any(word in context for word in ["required", "must", "essential", "minimum", "mandatory"]):
                        confidence = 0.9
                    elif any(word in context for word in ["preferred", "nice to have", "bonus", "plus"]):
                        confidence = 0.6
                    elif any(word in context for word in ["experience with", "knowledge of", "familiarity"]):
                        confidence = 0.8
                    
                    extracted_skills.append({
                        "skill": skill_name.title(),
                        "category": category,
                        "confidence": confidence,
                        "context": context.strip()
                    })
        
        # Remove duplicates and sort by confidence
        unique_skills = {}
        for skill in extracted_skills:
            skill_name = skill["skill"].lower()
            if skill_name not in unique_skills or skill["confidence"] > unique_skills[skill_name]["confidence"]:
                unique_skills[skill_name] = skill
        
        return sorted(unique_skills.values(), key=lambda x: x["confidence"], reverse=True)
    
    def _extract_soft_skills(self, text: str) -> List[Dict[str, Any]]:
        """Extract soft skills from text."""
        
        soft_skill_patterns = [
            r"communication", r"leadership", r"teamwork", r"problem solving", r"critical thinking",
            r"time management", r"organization", r"adaptability", r"creativity", r"attention to detail",
            r"analytical", r"collaboration", r"mentoring", r"project management", r"customer service"
        ]
        
        extracted_skills = []
        
        for pattern in soft_skill_patterns:
            matches = re.finditer(pattern, text.lower())
            for match in matches:
                skill_name = match.group()
                context_start = max(0, match.start() - 50)
                context_end = min(len(text), match.end() + 50)
                context = text[context_start:context_end].lower()
                
                confidence = 0.7
                if any(word in context for word in ["required", "must", "essential"]):
                    confidence = 0.9
                
                extracted_skills.append({
                    "skill": skill_name.title(),
                    "category": "soft_skills",
                    "confidence": confidence,
                    "context": context.strip()
                })
        
        return extracted_skills
    
    def _extract_experience_requirements(self, text: str) -> Dict[str, Any]:
        """Extract experience requirements from text."""
        
        # Look for experience patterns
        experience_patterns = [
            r"(\d+)\+?\s*years?\s*of\s*experience",
            r"(\d+)\+?\s*years?\s*in\s*the\s*field",
            r"minimum\s*(\d+)\s*years",
            r"at\s*least\s*(\d+)\s*years"
        ]
        
        for pattern in experience_patterns:
            match = re.search(pattern, text.lower())
            if match:
                years = int(match.group(1))
                return {
                    "minimum_years": years,
                    "pattern_found": match.group(0),
                    "confidence": 0.8
                }
        
        return {
            "minimum_years": None,
            "pattern_found": None,
            "confidence": 0.0
        }
    
    def _extract_education_requirements(self, text: str) -> Dict[str, Any]:
        """Extract education requirements from text."""
        
        education_keywords = [
            "bachelor", "master", "phd", "degree", "diploma", "certification", "certificate"
        ]
        
        requirements = []
        for keyword in education_keywords:
            if keyword in text.lower():
                requirements.append(keyword)
        
        return {
            "required_education": requirements,
            "confidence": 0.7 if requirements else 0.0
        }
    
    def _extract_keywords(self, text: str) -> List[str]:
        """Extract important keywords from job description."""
        
        # Remove common words and extract meaningful terms
        common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
        
        # Extract words with 3+ characters
        words = re.findall(r'\b\w{3,}\b', text.lower())
        
        # Filter out common words and count frequency
        from collections import Counter
        word_counts = Counter([word for word in words if word not in common_words])
        
        # Return top keywords
        return [word for word, count in word_counts.most_common(20)]
    
    def _analyze_current_resume(self, resume_content: str) -> Dict[str, Any]:
        """Analyze current resume content."""
        
        # Extract sections
        sections = self._extract_resume_sections(resume_content)
        
        # Extract skills
        skills = self._extract_resume_skills(resume_content)
        
        # Extract experience
        experience = self._extract_resume_experience(resume_content)
        
        # Extract education
        education = self._extract_resume_education(resume_content)
        
        # Calculate statistics
        word_count = len(resume_content.split())
        skill_count = len(skills)
        
        return {
            "sections": sections,
            "skills": skills,
            "experience": experience,
            "education": education,
            "statistics": {
                "word_count": word_count,
                "skill_count": skill_count,
                "section_count": len(sections)
            },
            "content": resume_content
        }
    
    def _extract_resume_sections(self, content: str) -> Dict[str, str]:
        """Extract sections from resume content."""
        
        # Common section headers
        section_headers = [
            "summary", "objective", "experience", "work experience", "employment",
            "education", "skills", "technical skills", "certifications", "projects",
            "achievements", "awards", "publications", "languages"
        ]
        
        sections = {}
        lines = content.split('\n')
        current_section = "header"
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            is_header = any(header in line_lower for header in section_headers)
            
            if is_header:
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                # Start new section
                current_section = line_lower
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    def _extract_resume_skills(self, content: str) -> List[Dict[str, Any]]:
        """Extract skills from resume content."""
        
        # Use the same skill patterns as job analysis
        skill_patterns = {
            "programming_languages": [
                r"python", r"java", r"javascript", r"typescript", r"c\+\+", r"c#", r"go", r"rust", r"php", r"ruby"
            ],
            "frameworks": [
                r"django", r"flask", r"react", r"angular", r"vue", r"node\.js", r"spring", r"laravel", r"rails"
            ],
            "databases": [
                r"postgresql", r"mysql", r"mongodb", r"redis", r"elasticsearch", r"dynamodb", r"sqlite"
            ],
            "cloud_platforms": [
                r"aws", r"azure", r"gcp", r"docker", r"kubernetes", r"terraform", r"ansible"
            ],
            "tools": [
                r"git", r"jenkins", r"jira", r"confluence", r"slack", r"figma", r"sketch"
            ]
        }
        
        skills = []
        
        for category, patterns in skill_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, content.lower())
                for match in matches:
                    skill_name = match.group()
                    context_start = max(0, match.start() - 30)
                    context_end = min(len(content), match.end() + 30)
                    context = content[context_start:context_end]
                    
                    skills.append({
                        "skill": skill_name.title(),
                        "category": category,
                        "context": context.strip(),
                        "mention_count": 1
                    })
        
        # Count mentions and remove duplicates
        skill_counts = {}
        for skill in skills:
            skill_name = skill["skill"].lower()
            if skill_name in skill_counts:
                skill_counts[skill_name]["mention_count"] += 1
            else:
                skill_counts[skill_name] = skill
        
        return list(skill_counts.values())
    
    def _extract_resume_experience(self, content: str) -> List[Dict[str, Any]]:
        """Extract work experience from resume content."""
        
        # Look for experience patterns
        experience_patterns = [
            r"(\d{4})\s*[-–]\s*(\d{4}|present|current)",
            r"(\d{4})\s*to\s*(\d{4}|present|current)"
        ]
        
        experiences = []
        
        for pattern in experience_patterns:
            matches = re.finditer(pattern, content)
            for match in matches:
                start_year = match.group(1)
                end_year = match.group(2)
                
                # Extract context around the date
                context_start = max(0, match.start() - 100)
                context_end = min(len(content), match.end() + 100)
                context = content[context_start:context_end]
                
                experiences.append({
                    "start_year": start_year,
                    "end_year": end_year,
                    "context": context.strip()
                })
        
        return experiences
    
    def _extract_resume_education(self, content: str) -> List[Dict[str, Any]]:
        """Extract education information from resume content."""
        
        education_keywords = ["bachelor", "master", "phd", "degree", "university", "college", "school"]
        
        education = []
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in education_keywords):
                # Extract education line and context
                context_start = max(0, i - 2)
                context_end = min(len(lines), i + 3)
                context = '\n'.join(lines[context_start:context_end])
                
                education.append({
                    "line": line.strip(),
                    "context": context.strip()
                })
        
        return education
    
    def _identify_skill_gaps(self, resume_analysis: Dict[str, Any], 
                            job_analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Identify gaps between current skills and job requirements."""
        
        current_skills = {skill["skill"].lower() for skill in resume_analysis.get("skills", [])}
        required_skills = job_analysis.get("technical_skills", [])
        
        missing_skills = []
        matching_skills = []
        skill_enhancements = []
        
        for req_skill in required_skills:
            skill_name = req_skill["skill"].lower()
            
            if skill_name in current_skills:
                # Find current skill details
                current_skill = next(
                    (s for s in resume_analysis["skills"] if s["skill"].lower() == skill_name), 
                    None
                )
                
                matching_skills.append({
                    "skill": req_skill["skill"],
                    "category": req_skill["category"],
                    "confidence": req_skill["confidence"],
                    "current_mentions": current_skill["mention_count"] if current_skill else 0
                })
                
                # Check if skill needs enhancement
                if current_skill and current_skill["mention_count"] < 2 and req_skill["confidence"] > 0.7:
                    skill_enhancements.append({
                        "skill": req_skill["skill"],
                        "category": req_skill["category"],
                        "current_mentions": current_skill["mention_count"],
                        "suggested_mentions": 3
                    })
            else:
                missing_skills.append({
                    "skill": req_skill["skill"],
                    "category": req_skill["category"],
                    "confidence": req_skill["confidence"],
                    "priority": "high" if req_skill["confidence"] > 0.8 else "medium"
                })
        
        return {
            "missing_skills": missing_skills,
            "matching_skills": matching_skills,
            "skill_enhancements": skill_enhancements,
            "coverage_percentage": len(matching_skills) / len(required_skills) * 100 if required_skills else 0,
            "total_required": len(required_skills),
            "total_matching": len(matching_skills),
            "total_missing": len(missing_skills)
        }
    
    async def _generate_optimized_content(self, resume_analysis: Dict[str, Any],
                                        job_analysis: Dict[str, Any],
                                        skill_gaps: Dict[str, Any],
                                        job_title: str) -> Dict[str, Any]:
        """Generate optimized resume content using AI."""
        
        if not self.client:
            return self._generate_basic_optimized_content(resume_analysis, job_analysis, skill_gaps)
        
        try:
            # Prepare prompt for AI
            prompt = self._create_optimization_prompt(resume_analysis, job_analysis, skill_gaps, job_title)
            
            response = await self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert resume writer and career coach. Create compelling, truthful, and ATS-optimized resume content."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.3
            )
            
            # Parse AI response
            ai_response = response.choices[0].message.content
            optimized_data = self._parse_ai_response(ai_response)
            
            return optimized_data
            
        except Exception as e:
            # Fallback to basic optimization
            return self._generate_basic_optimized_content(resume_analysis, job_analysis, skill_gaps)
    
    def _create_optimization_prompt(self, resume_analysis: Dict[str, Any],
                                  job_analysis: Dict[str, Any],
                                  skill_gaps: Dict[str, Any],
                                  job_title: str) -> str:
        """Create prompt for AI optimization."""
        
        current_skills = [skill["skill"] for skill in resume_analysis.get("skills", [])]
        missing_skills = [skill["skill"] for skill in skill_gaps.get("missing_skills", [])]
        required_skills = [skill["skill"] for skill in job_analysis.get("technical_skills", [])]
        
        prompt = f"""
        Optimize this resume for a {job_title} position. The goal is to better align the resume with the job requirements while maintaining truthfulness and professionalism.
        
        Current Resume Content:
        {resume_analysis.get("content", "")[:2000]}
        
        Job Requirements:
        - Technical Skills: {', '.join(required_skills[:15])}
        - Missing Skills: {', '.join(missing_skills[:10])}
        - Current Skills: {', '.join(current_skills[:15])}
        
        Please provide optimized content in JSON format:
        {{
            "optimized_summary": "Professional summary optimized for this role...",
            "enhanced_experience": [
                {{
                    "title": "Job Title",
                    "company": "Company Name", 
                    "description": "Enhanced job description with relevant keywords...",
                    "achievements": ["Achievement 1", "Achievement 2"]
                }}
            ],
            "skills_section": {{
                "technical_skills": ["skill1", "skill2", ...],
                "core_competencies": ["competency1", "competency2", ...]
            }},
            "keywords_added": ["keyword1", "keyword2", ...],
            "modifications_summary": ["modification1", "modification2", ...],
            "ats_optimization_tips": ["tip1", "tip2", ...]
        }}
        
        Guidelines:
        1. Keep all information truthful - only enhance presentation, don't fabricate
        2. Incorporate high-priority required skills naturally into existing experience
        3. Use action verbs and quantifiable achievements
        4. Optimize for ATS (Applicant Tracking Systems)
        5. Maintain professional tone and format
        6. Focus on relevant experience and skills
        """
        
        return prompt
    
    def _parse_ai_response(self, ai_response: str) -> Dict[str, Any]:
        """Parse AI response to extract JSON content."""
        
        try:
            # Extract JSON from the response
            json_start = ai_response.find('{')
            json_end = ai_response.rfind('}') + 1
            
            if json_start != -1 and json_end != -1:
                json_str = ai_response[json_start:json_end]
                optimized_data = json.loads(json_str)
                
                # Add metadata
                optimized_data["source"] = "ai"
                optimized_data["optimization_timestamp"] = datetime.now().isoformat()
                
                return optimized_data
            
        except json.JSONDecodeError:
            pass
        
        # Return basic structure if parsing fails
        return {
            "source": "fallback",
            "optimization_timestamp": datetime.now().isoformat(),
            "optimized_summary": "Resume optimization completed",
            "enhanced_experience": [],
            "skills_section": {"technical_skills": [], "core_competencies": []},
            "keywords_added": [],
            "modifications_summary": ["Basic optimization applied"],
            "ats_optimization_tips": ["Use relevant keywords", "Maintain clean format"]
        }
    
    def _generate_basic_optimized_content(self, resume_analysis: Dict[str, Any],
                                        job_analysis: Dict[str, Any],
                                        skill_gaps: Dict[str, Any]) -> Dict[str, Any]:
        """Generate basic optimized content without AI."""
        
        current_skills = [skill["skill"] for skill in resume_analysis.get("skills", [])]
        missing_skills = [skill["skill"] for skill in skill_gaps.get("missing_skills", [])]
        
        # Create basic optimization
        optimized_summary = resume_analysis.get("sections", {}).get("summary", "")
        if missing_skills:
            optimized_summary += f" Skilled in {', '.join(missing_skills[:3])} and related technologies."
        
        return {
            "optimized_summary": optimized_summary,
            "enhanced_experience": resume_analysis.get("experience", []),
            "skills_section": {
                "technical_skills": current_skills + missing_skills[:5],
                "core_competencies": ["Problem Solving", "Team Collaboration", "Continuous Learning"]
            },
            "keywords_added": missing_skills[:5],
            "modifications_summary": [f"Added {len(missing_skills)} missing skills", "Enhanced summary"],
            "ats_optimization_tips": ["Include relevant keywords", "Use action verbs", "Quantify achievements"],
            "source": "basic",
            "optimization_timestamp": datetime.now().isoformat()
        }
    
    def _create_modification_plan(self, resume_analysis: Dict[str, Any],
                                job_analysis: Dict[str, Any],
                                skill_gaps: Dict[str, Any],
                                optimized_content: Dict[str, Any]) -> Dict[str, Any]:
        """Create a detailed plan for resume modifications."""
        
        plan = {
            "summary": "Resume modification plan to better align with job requirements",
            "priority_changes": [],
            "skill_additions": [],
            "content_enhancements": [],
            "ats_optimizations": [],
            "estimated_time": "15-30 minutes"
        }
        
        # Priority changes based on skill gaps
        missing_skills = skill_gaps.get("missing_skills", [])
        high_priority_missing = [skill for skill in missing_skills if skill.get("priority") == "high"]
        
        if high_priority_missing:
            plan["priority_changes"].append({
                "type": "add_skills",
                "description": f"Add {len(high_priority_missing)} high-priority missing skills",
                "skills": [skill["skill"] for skill in high_priority_missing],
                "priority": "high"
            })
        
        # Content enhancements
        if optimized_content.get("optimized_summary"):
            plan["content_enhancements"].append({
                "type": "update_summary",
                "description": "Update professional summary to highlight relevant experience",
                "priority": "high"
            })
        
        # ATS optimizations
        plan["ats_optimizations"].extend([
            {
                "type": "keyword_optimization",
                "description": "Incorporate job-specific keywords naturally",
                "priority": "medium"
            },
            {
                "type": "format_optimization", 
                "description": "Ensure clean, scannable format for ATS systems",
                "priority": "medium"
            }
        ])
        
        return plan
    
    def _generate_ats_recommendations(self, resume_analysis: Dict[str, Any],
                                    job_analysis: Dict[str, Any],
                                    skill_gaps: Dict[str, Any]) -> List[str]:
        """Generate ATS optimization recommendations."""
        
        recommendations = []
        
        # Keyword optimization
        missing_skills = skill_gaps.get("missing_skills", [])
        if missing_skills:
            top_missing = [skill["skill"] for skill in missing_skills[:3]]
            recommendations.append(f"Incorporate these keywords naturally: {', '.join(top_missing)}")
        
        # Format recommendations
        recommendations.extend([
            "Use standard section headings (Experience, Education, Skills)",
            "Avoid tables, graphics, or complex formatting",
            "Use standard fonts (Arial, Calibri, Times New Roman)",
            "Keep file size under 2MB",
            "Use bullet points for achievements and responsibilities"
        ])
        
        # Content recommendations
        recommendations.extend([
            "Quantify achievements with numbers and percentages",
            "Use action verbs at the beginning of bullet points",
            "Include relevant certifications and training",
            "Tailor experience descriptions to match job requirements"
        ])
        
        return recommendations
    
    async def _create_modified_resume(self, original_resume_path: str,
                                    optimized_content: Dict[str, Any],
                                    job_title: str,
                                    company_name: str,
                                    output_dir: Optional[str] = None) -> str:
        """Create a modified resume file based on the optimization results."""
        
        try:
            # Create output directory
            if not output_dir:
                output_dir = os.path.join(os.path.dirname(original_resume_path), "modified_resumes")
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_job_title = re.sub(r'[^\w\s-]', '', job_title).replace(' ', '_')[:30]
            safe_company = re.sub(r'[^\w\s-]', '', company_name).replace(' ', '_')[:20]
            
            filename = f"resume_{safe_job_title}_{safe_company}_{timestamp}.docx"
            output_path = os.path.join(output_dir, filename)
            
            # Create the modified document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'Resume - {job_title}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add modification info
            doc.add_paragraph(f'Modified for: {job_title} at {company_name}')
            doc.add_paragraph(f'Modification date: {datetime.now().strftime("%B %d, %Y")}')
            doc.add_paragraph('')
            
            # Add optimized content
            # Summary section
            if optimized_content.get("optimized_summary"):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(optimized_content["optimized_summary"])
                doc.add_paragraph('')
            
            # Skills section
            skills_section = optimized_content.get("skills_section", {})
            if skills_section:
                doc.add_heading('Technical Skills', level=1)
                
                if skills_section.get("technical_skills"):
                    doc.add_paragraph(', '.join(skills_section["technical_skills"]))
                
                if skills_section.get("core_competencies"):
                    doc.add_paragraph('')
                    doc.add_paragraph('Core Competencies:')
                    doc.add_paragraph(', '.join(skills_section["core_competencies"]))
                
                doc.add_paragraph('')
            
            # Experience section (enhanced)
            enhanced_experience = optimized_content.get("enhanced_experience", [])
            if enhanced_experience:
                doc.add_heading('Professional Experience', level=1)
                
                for exp in enhanced_experience:
                    # Job title and company
                    job_header = doc.add_heading(f'{exp.get("title", "Job Title")} - {exp.get("company", "Company")}', level=2)
                    
                    # Description
                    if exp.get("description"):
                        doc.add_paragraph(exp["description"])
                    
                    # Achievements
                    if exp.get("achievements"):
                        for achievement in exp["achievements"]:
                            p = doc.add_paragraph()
                            p.add_run('• ').bold = True
                            p.add_run(achievement)
                    
                    doc.add_paragraph('')
            
            # Add modification summary
            modifications = optimized_content.get("modifications_summary", [])
            if modifications:
                doc.add_heading('Modifications Made', level=1)
                for mod in modifications:
                    doc.add_paragraph(f'• {mod}')
                doc.add_paragraph('')
            
            # Add ATS recommendations
            ats_recs = optimized_content.get("ats_optimization_tips", [])
            if ats_recs:
                doc.add_heading('ATS Optimization Tips', level=1)
                for rec in ats_recs[:5]:  # Top 5 recommendations
                    doc.add_paragraph(f'• {rec}')
            
            # Save the document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Failed to create modified resume: {str(e)}")
    
    async def get_editing_tools(self) -> List[Dict[str, Any]]:
        """Get available resume editing tools."""
        
        return [
            {
                "name": "edit_resume_for_job",
                "description": "Edit resume to match specific job requirements",
                "parameters": {
                    "resume_path": "Path to original resume",
                    "job_title": "Target job title",
                    "job_description": "Job description text",
                    "company_name": "Company name",
                    "output_dir": "Output directory (optional)"
                }
            },
            {
                "name": "analyze_job_requirements",
                "description": "Analyze job description to extract requirements",
                "parameters": {
                    "job_description": "Job description text"
                }
            },
            {
                "name": "analyze_resume",
                "description": "Analyze current resume content",
                "parameters": {
                    "resume_path": "Path to resume file"
                }
            },
            {
                "name": "identify_skill_gaps",
                "description": "Identify gaps between current skills and job requirements",
                "parameters": {
                    "resume_analysis": "Resume analysis data",
                    "job_analysis": "Job analysis data"
                }
            }
        ]

