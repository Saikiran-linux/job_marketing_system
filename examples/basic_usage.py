#!/usr/bin/env python3
"""
Basic usage examples for the Multi-Agent Job Application System.
"""

import asyncio
import os
import sys

# Add project root to path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from agents.orchestrator_agent import OrchestratorAgent
from agents.job_search_agent import JobSearchAgent
from agents.skills_analysis_agent import SkillsAnalysisAgent
from agents.resume_analysis_agent import ResumeAnalysisAgent
from utils.report_generator import ReportGenerator

async def example_1_complete_workflow():
    """Example 1: Complete job search and analysis workflow."""
    
    print("=" * 60)
    print("EXAMPLE 1: Complete Workflow")
    print("=" * 60)
    
    orchestrator = OrchestratorAgent()
    
    # Note: You'll need to provide a real resume file path
    resume_path = "./data/sample_resume.docx"
    
    # Create a sample resume if it doesn't exist
    if not os.path.exists(resume_path):
        os.makedirs(os.path.dirname(resume_path), exist_ok=True)
        create_sample_resume(resume_path)
    
    from agents.base_agent import AgentState
    from datetime import datetime
    
    # Create proper AgentState object
    initial_state = AgentState(
        session_id="",
        start_time="",
        current_step="",
        steps_completed=[],
        status="",
        error=None,
        role="Software Engineer",
        resume_path=resume_path,
        location="remote",
        max_jobs=5,  # Limit for demo
        auto_apply=False,  # Safe mode - no actual applications
        resume_analysis=None,
        job_search_results=None,
        processed_jobs=[],
        final_report=None,
        end_time=None,
        workflow_duration=None
    )
    
    result = await orchestrator.execute(initial_state)
    
    if result.status == "completed":
        print(f"✅ Workflow completed successfully!")
        print(f"   Session ID: {result.session_id}")
        print(f"   Steps completed: {len(result.steps_completed)}")
        
        # Show some processed jobs
        if result.processed_jobs:
            print(f"\n📋 Sample processed jobs:")
            for i, job in enumerate(result.processed_jobs[:3], 1):
                print(f"   {i}. {job.get('title', 'Unknown')} at {job.get('company', 'Unknown')}")
    else:
        print(f"❌ Workflow failed: {result.error or 'Unknown error'}")

async def example_2_job_search_only():
    """Example 2: Job search without applications."""
    
    print("\n" + "=" * 60)
    print("EXAMPLE 2: Job Search Only")
    print("=" * 60)
    
    from agents.base_agent import AgentState
    job_search_agent = JobSearchAgent()
    
    # Create proper AgentState object for job search
    job_search_state = AgentState(
        session_id="",
        start_time="",
        current_step="",
        steps_completed=[],
        status="",
        error=None,
        role="Python Developer",
        resume_path="./sample_resume.docx",
        location="San Francisco",
        max_jobs=10,
        auto_apply=False,
        resume_analysis=None,
        job_search_results=None,
        processed_jobs=[],
        final_report=None,
        end_time=None,
        workflow_duration=None
    )
    
    result = await job_search_agent.execute(job_search_state)
    
    if result.status == "completed":
        if result.job_search_results and result.job_search_results.get("jobs"):
            jobs = result.job_search_results["jobs"]
            print(f"✅ Found {len(jobs)} jobs")
            
            for i, job in enumerate(jobs[:5], 1):  # Show first 5
                print(f"   {i}. {job.get('title', 'Unknown')} at {job.get('company', 'Unknown')}")
                print(f"      Location: {job.get('location', 'Unknown')}")
                print(f"      Source: {job.get('source', 'Unknown')}")
                print()
        else:
            print("✅ Job search completed but no jobs found")
    else:
        print(f"❌ Job search failed: {result.error or 'Unknown error'}")

async def example_3_skills_analysis():
    """Example 3: Analyze skills from a job description."""
    
    print("=" * 60)
    print("EXAMPLE 3: Skills Analysis")
    print("=" * 60)
    
    from agents.base_agent import AgentState
    
    # Sample job description
    job_description = """
    We are looking for a Senior Python Developer to join our team.
    
    Required Skills:
    - 5+ years of Python development experience
    - Strong knowledge of Django and Flask frameworks
    - Experience with PostgreSQL and MongoDB
    - Familiarity with AWS services (EC2, S3, RDS)
    - Knowledge of Docker and Kubernetes
    - Experience with Git version control
    - Strong problem-solving skills
    - Excellent communication abilities
    
    Preferred Skills:
    - React.js for frontend development
    - Machine learning experience with scikit-learn or TensorFlow
    - Experience with CI/CD pipelines
    - Knowledge of microservices architecture
    """
    
    skills_agent = SkillsAnalysisAgent()
    
    # Create proper AgentState object for skills analysis
    skills_state = AgentState(
        session_id="",
        start_time="",
        current_step="",
        steps_completed=[],
        status="",
        error=None,
        role="Senior Python Developer",
        resume_path="./sample_resume.docx",
        location="",
        max_jobs=1,
        auto_apply=False,
        resume_analysis=None,
        job_search_results=None,
        processed_jobs=[],
        final_report=None,
        end_time=None,
        workflow_duration=None
    )
    
    # For skills analysis, we need to pass job description in a different way
    # Let's modify the state to include job description
    skills_state.job_search_results = {
        "job_description": job_description,
        "job_title": "Senior Python Developer"
    }
    
    result = await skills_agent.execute(skills_state)
    
    if result.status == "completed":
        # Extract skills from the processed state
        if result.processed_jobs and len(result.processed_jobs) > 0:
            job_result = result.processed_jobs[0]
            required_skills = job_result.get("required_skills", [])
            skill_categories = job_result.get("skill_categories", {})
            
            print(f"✅ Extracted {len(required_skills)} skills")
            
            print(f"\n🔍 Top skills found:")
        for skill in required_skills[:10]:  # Top 10 skills
            skill_name = skill.get("skill", "Unknown")
            confidence = skill.get("confidence", 0)
            print(f"   • {skill_name} (confidence: {confidence:.2f})")
        
        print(f"\n📊 Skill categories:")
        for category, skills in skill_categories.items():
            if skills:
                print(f"   {category.title()}: {len(skills)} skills")
    else:
        print(f"❌ Skills analysis failed: {result.error or 'Unknown error'}")

async def example_4_resume_analysis():
    """Example 4: Analyze an existing resume."""
    
    print("\n" + "=" * 60)
    print("EXAMPLE 4: Resume Analysis")
    print("=" * 60)
    
    resume_path = "./data/sample_resume.docx"
    
    # Create sample resume if needed
    if not os.path.exists(resume_path):
        os.makedirs(os.path.dirname(resume_path), exist_ok=True)
        create_sample_resume(resume_path)
    
    resume_agent = ResumeAnalysisAgent()
    
    result = await resume_agent.execute({
        "resume_path": resume_path
    })
    
    if result.get("status") == "success":
        current_skills = result.get("current_skills", [])
        sections = result.get("sections", {})
        stats = result.get("statistics", {})
        
        print(f"✅ Resume analyzed successfully")
        print(f"   Total words: {stats.get('total_words', 0)}")
        print(f"   Sections found: {len(sections)}")
        print(f"   Skills extracted: {len(current_skills)}")
        
        print(f"\n📋 Resume sections:")
        for section_name in sections.keys():
            print(f"   • {section_name.title()}")
        
        print(f"\n🛠️ Top skills found:")
        for skill in current_skills[:10]:  # Top 10 skills
            skill_name = skill.get("skill", "Unknown")
            mentions = skill.get("mention_count", 0)
            print(f"   • {skill_name} (mentioned {mentions} times)")
    else:
        print(f"❌ Resume analysis failed: {result.get('message', 'Unknown error')}")

def example_5_generate_reports():
    """Example 5: Generate analytics reports."""
    
    print("\n" + "=" * 60)
    print("EXAMPLE 5: Report Generation")
    print("=" * 60)
    
    reporter = ReportGenerator()
    
    # Generate weekly report
    try:
        weekly_report = reporter.generate_weekly_report()
        
        print("✅ Weekly report generated")
        print(f"   Period: {weekly_report['period']['start_date']} to {weekly_report['period']['end_date']}")
        
        stats = weekly_report.get("statistics", {})
        print(f"   Total applications: {stats.get('total_applications', 0)}")
        
        insights = weekly_report.get("insights", [])
        if insights:
            print(f"\n💡 Insights:")
            for insight in insights[:3]:
                print(f"   • {insight}")
    
    except Exception as e:
        print(f"⚠️ Report generation failed: {str(e)}")
        print("   This is normal if no application data exists yet")
    
    # Generate skill gap report
    try:
        skill_report = reporter.generate_skill_gap_report("Software Engineer")
        
        print(f"\n✅ Skill gap report generated")
        recommendations = skill_report.get("recommendations", [])
        if recommendations:
            print(f"   Recommendations: {len(recommendations)}")
            for rec in recommendations[:3]:
                print(f"   • {rec.get('recommendation', 'No recommendation')}")
    
    except Exception as e:
        print(f"⚠️ Skill gap report failed: {str(e)}")

def create_sample_resume(file_path):
    """Create a sample resume for demonstration."""
    
    try:
        from docx import Document
        
        doc = Document()
        
        # Add title
        doc.add_heading('John Doe', 0)
        
        # Contact info
        doc.add_paragraph('Email: john.doe@example.com | Phone: (555) 123-4567')
        doc.add_paragraph('LinkedIn: linkedin.com/in/johndoe | GitHub: github.com/johndoe')
        
        # Professional summary
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(
            'Experienced Software Engineer with 5+ years in Python development. '
            'Skilled in web frameworks, databases, and cloud technologies. '
            'Passionate about creating scalable solutions and continuous learning.'
        )
        
        # Skills
        doc.add_heading('Technical Skills', level=1)
        doc.add_paragraph(
            'Programming Languages: Python, JavaScript, TypeScript, Java\n'
            'Frameworks: Django, Flask, React, Node.js\n'
            'Databases: PostgreSQL, MongoDB, Redis\n'
            'Cloud: AWS (EC2, S3, RDS), Docker, Kubernetes\n'
            'Tools: Git, Jenkins, Jira, VS Code'
        )
        
        # Experience
        doc.add_heading('Professional Experience', level=1)
        
        doc.add_heading('Senior Python Developer - Tech Corp (2021-Present)', level=2)
        doc.add_paragraph('• Developed microservices using Django and PostgreSQL')
        doc.add_paragraph('• Implemented CI/CD pipelines reducing deployment time by 50%')
        doc.add_paragraph('• Led team of 3 developers on customer-facing web applications')
        
        doc.add_heading('Python Developer - StartupXYZ (2019-2021)', level=2)
        doc.add_paragraph('• Built REST APIs using Flask and SQLAlchemy')
        doc.add_paragraph('• Optimized database queries improving performance by 30%')
        doc.add_paragraph('• Collaborated with frontend team on React integration')
        
        # Education
        doc.add_heading('Education', level=1)
        doc.add_paragraph('Bachelor of Science in Computer Science - University ABC (2019)')
        
        doc.save(file_path)
        print(f"✅ Sample resume created at: {file_path}")
        
    except ImportError:
        print("⚠️ python-docx not available, cannot create sample resume")
    except Exception as e:
        print(f"⚠️ Error creating sample resume: {str(e)}")

async def main():
    """Run all examples."""
    
    print("🤖 Multi-Agent Job Application System - Examples")
    print("=" * 60)
    print("This demo shows the system capabilities without making real applications.")
    print("To enable actual job applications, set auto_apply=True (use with caution!)")
    print()
    
    try:
        # Run examples
        await example_1_complete_workflow()
        await example_2_job_search_only()
        await example_3_skills_analysis()
        await example_4_resume_analysis()
        example_5_generate_reports()
        
        print("\n" + "=" * 60)
        print("✅ All examples completed!")
        print("=" * 60)
        print("\nNext steps:")
        print("1. Set up your .env file with API keys")
        print("2. Add your real resume file")
        print("3. Run: python main.py --role 'Your Target Role' --resume 'your_resume.docx'")
        print("4. Use --auto-apply flag when you're ready for real applications")
        
    except KeyboardInterrupt:
        print("\n❌ Examples interrupted by user")
    except Exception as e:
        print(f"\n❌ Example failed: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    asyncio.run(main())
