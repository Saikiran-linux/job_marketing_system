import asyncio
import os
import re
import json
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import openai
from openai import AsyncOpenAI
from agents.base_agent import BaseAgent, AgentState
from config import Config

class ResumeModificationAgent(BaseAgent):
    """Agent responsible for modifying resumes to match job requirements."""
    
    def __init__(self):
        super().__init__("ResumeModificationAgent")
        self.client = AsyncOpenAI(api_key=Config.OPENAI_API_KEY) if Config.OPENAI_API_KEY else None
    
    async def execute(self, state: AgentState) -> AgentState:
        """Modify resume to match job requirements."""
        
        try:
            # Check if we have job information in the state
            if not state.job_search_results:
                state.error = "No job information available for resume modification"
                return state
            
            # Get job details
            job_info = state.job_search_results.get("current_job", {})
            if not job_info:
                state.error = "No current job information for resume modification"
                return state
            
            job_title = job_info.get("title", "")
            job_description = job_info.get("description", "")
            company_name = job_info.get("company", "")
            
            # Get resume analysis
            if not state.resume_analysis:
                state.error = "No resume analysis available for modification"
                return state
            
            # Perform resume modification
            modification_result = await self.modify_resume_for_job(
                state.resume_analysis,
                job_title,
                job_description,
                company_name,
                state.resume_path
            )
            
            if modification_result.get("status") == "success":
                # Update state with modification results
                state.resume_modification = modification_result
                state.steps_completed.append("resume_modification")
                state.current_step = "resume_modification_complete"
                
                # Create the modified resume file
                modified_resume_path = await self.create_modified_resume(
                    modification_result,
                    state.resume_path,
                    job_title,
                    company_name
                )
                
                if modified_resume_path:
                    state.resume_modification["modified_resume_path"] = modified_resume_path
                    self.log_action("SUCCESS", f"Resume modified and saved to: {modified_resume_path}")
                else:
                    self.log_action("WARNING", "Resume modification completed but file creation failed")
                
            else:
                state.error = f"Resume modification failed: {modification_result.get('error', 'Unknown error')}"
            
            return state
            
        except Exception as e:
            self.log_action("ERROR", f"Resume modification failed: {str(e)}")
            state.error = f"Resume modification failed: {str(e)}"
            return state
    
    async def modify_resume_for_job(self, resume_analysis: Dict[str, Any], 
                                  job_title: str, job_description: str, 
                                  company_name: str, original_resume_path: str) -> Dict[str, Any]:
        """Modify resume to match specific job requirements."""
        
        try:
            self.log_action("MODIFYING", f"Modifying resume for {job_title} at {company_name}")
            
            # Analyze skill gaps
            skill_gaps = self._analyze_skill_gaps(resume_analysis, job_description)
            
            # Generate optimized content
            optimized_content = await self._generate_optimized_content(
                resume_analysis, skill_gaps, job_description, job_title
            )
            
            # Create modification plan
            modification_plan = self._create_modification_plan(
                resume_analysis, skill_gaps, optimized_content, job_title
            )
            
            # Generate ATS optimization recommendations
            ats_recommendations = self._generate_ats_recommendations(
                resume_analysis, job_description, skill_gaps
            )
            
            return {
                "status": "success",
                "job_title": job_title,
                "company_name": company_name,
                "skill_gaps": skill_gaps,
                "optimized_content": optimized_content,
                "modification_plan": modification_plan,
                "ats_recommendations": ats_recommendations,
                "modification_timestamp": datetime.now().isoformat(),
                "original_resume_path": original_resume_path
            }
            
        except Exception as e:
            self.log_action("ERROR", f"Resume modification failed: {str(e)}")
            return {
                "status": "error",
                "error": str(e),
                "modification_timestamp": datetime.now().isoformat()
            }
    
    def _analyze_skill_gaps(self, resume_analysis: Dict[str, Any], 
                           job_description: str) -> Dict[str, Any]:
        """Analyze gaps between current skills and required skills from job description."""
        
        # Extract skills from resume
        current_skills = resume_analysis.get("current_skills", [])
        current_skill_names = {skill["skill"].lower() for skill in current_skills}
        
        # Extract required skills from job description using AI or pattern matching
        required_skills = self._extract_required_skills(job_description)
        
        # Categorize skills
        missing_skills = []
        matching_skills = []
        skill_enhancements = []
        
        for req_skill in required_skills:
            skill_name = req_skill.get("skill", "").lower()
            skill_confidence = req_skill.get("confidence", 0)
            
            if skill_name in current_skill_names:
                # Find current skill details
                current_skill = next(
                    (s for s in current_skills if s["skill"].lower() == skill_name), 
                    None
                )
                
                matching_skills.append({
                    "skill": skill_name,
                    "current_mentions": current_skill["mention_count"] if current_skill else 0,
                    "required_confidence": skill_confidence,
                    "enhancement_needed": skill_confidence > 0.8 and 
                                        (current_skill["mention_count"] if current_skill else 0) < 2
                })
                
                # Check if skill needs enhancement
                if current_skill and current_skill["mention_count"] < 2 and skill_confidence > 0.7:
                    skill_enhancements.append({
                        "skill": skill_name,
                        "current_mentions": current_skill["mention_count"],
                        "suggested_mentions": 3,
                        "contexts_to_add": req_skill.get("categories", [])
                    })
            else:
                missing_skills.append({
                    "skill": skill_name,
                    "confidence": skill_confidence,
                    "category": req_skill.get("categories", ["unknown"])[0] if req_skill.get("categories") else "unknown",
                    "priority": "high" if skill_confidence > 0.8 else "medium" if skill_confidence > 0.5 else "low"
                })
        
        return {
            "missing_skills": missing_skills,
            "matching_skills": matching_skills,
            "skill_enhancements": skill_enhancements,
            "coverage_percentage": len(matching_skills) / len(required_skills) * 100 if required_skills else 0,
            "total_required": len(required_skills),
            "total_matching": len(matching_skills),
            "total_missing": len(missing_skills)
        }
    
    def _extract_required_skills(self, job_description: str) -> List[Dict[str, Any]]:
        """Extract required skills from job description."""
        
        # Common technical skills patterns
        skill_patterns = {
            "programming_languages": [
                r"python", r"java", r"javascript", r"typescript", r"c\+\+", r"c#", r"go", r"rust", r"php", r"ruby"
            ],
            "frameworks": [
                r"django", r"flask", r"react", r"angular", r"vue", r"node\.js", r"spring", r"laravel", r"rails"
            ],
            "databases": [
                r"postgresql", r"mysql", r"mongodb", r"redis", r"elasticsearch", r"dynamodb", r"sqlite"
            ],
            "cloud_platforms": [
                r"aws", r"azure", r"gcp", r"docker", r"kubernetes", r"terraform", r"ansible"
            ],
            "tools": [
                r"git", r"jenkins", r"jira", r"confluence", r"slack", r"figma", r"sketch"
            ]
        }
        
        extracted_skills = []
        
        for category, patterns in skill_patterns.items():
            for pattern in patterns:
                matches = re.finditer(pattern, job_description.lower())
                for match in matches:
                    skill_name = match.group()
                    # Calculate confidence based on context
                    context_start = max(0, match.start() - 50)
                    context_end = min(len(job_description), match.end() + 50)
                    context = job_description[context_start:context_end].lower()
                    
                    # Higher confidence if skill is mentioned in requirements section
                    confidence = 0.7
                    if any(word in context for word in ["required", "must", "essential", "minimum"]):
                        confidence = 0.9
                    elif any(word in context for word in ["preferred", "nice to have", "bonus"]):
                        confidence = 0.6
                    
                    extracted_skills.append({
                        "skill": skill_name.title(),
                        "confidence": confidence,
                        "categories": [category],
                        "context": context.strip()
                    })
        
        # Remove duplicates and sort by confidence
        unique_skills = {}
        for skill in extracted_skills:
            skill_name = skill["skill"].lower()
            if skill_name not in unique_skills or skill["confidence"] > unique_skills[skill_name]["confidence"]:
                unique_skills[skill_name] = skill
        
        return sorted(unique_skills.values(), key=lambda x: x["confidence"], reverse=True)
    
    async def _generate_optimized_content(self, resume_analysis: Dict[str, Any], 
                                        skill_gaps: Dict[str, Any],
                                        job_description: str, 
                                        job_title: str) -> Dict[str, Any]:
        """Generate optimized resume content using AI."""
        
        if not self.client:
            return self._generate_basic_optimized_content(resume_analysis, skill_gaps)
        
        try:
            # Extract key information from current resume
            resume_content = resume_analysis.get("resume_content", "")
            current_skills = [skill["skill"] for skill in resume_analysis.get("current_skills", [])]
            experience_details = resume_analysis.get("experience_details", [])
            
            # Create required skills list
            required_skills = skill_gaps.get("missing_skills", []) + skill_gaps.get("matching_skills", [])
            required_skill_names = [skill.get("skill", "") for skill in required_skills]
            high_priority_skills = [
                skill.get("skill", "") for skill in required_skills 
                if skill.get("confidence", 0) > 0.7
            ]
            
            prompt = f"""
            Optimize this resume for a {job_title} position. The goal is to better align the resume with the job requirements while maintaining truthfulness and professionalism.
            
            Current Resume Content:
            {resume_content[:2000]}
            
            Job Description:
            {job_description[:1000]}
            
            Required Skills: {', '.join(required_skill_names)}
            High Priority Skills: {', '.join(high_priority_skills)}
            Current Skills: {', '.join(current_skills)}
            
            Please provide optimized content in JSON format:
            {{
                "optimized_summary": "Professional summary optimized for this role...",
                "enhanced_experience": [
                    {{
                        "title": "Job Title",
                        "company": "Company Name", 
                        "description": "Enhanced job description with relevant keywords...",
                        "achievements": ["Achievement 1", "Achievement 2"]
                    }}
                ],
                "skills_section": {{
                    "technical_skills": ["skill1", "skill2", ...],
                    "core_competencies": ["competency1", "competency2", ...]
                }},
                "keywords_added": ["keyword1", "keyword2", ...],
                "modifications_summary": ["modification1", "modification2", ...],
                "ats_optimization_tips": ["tip1", "tip2", ...]
            }}
            
            Guidelines:
            1. Keep all information truthful - only enhance presentation, don't fabricate
            2. Incorporate high-priority required skills naturally into existing experience
            3. Use action verbs and quantifiable achievements
            4. Optimize for ATS (Applicant Tracking Systems)
            5. Maintain professional tone and format
            6. Focus on relevant experience and skills
            """
            
            response = await self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert resume writer and career coach. Create compelling, truthful, and ATS-optimized resume content."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.3
            )
            
            # Parse the AI response
            ai_response = response.choices[0].message.content
            
            # Extract JSON from the response
            try:
                json_start = ai_response.find('{')
                json_end = ai_response.rfind('}') + 1
                
                if json_start != -1 and json_end != -1:
                    json_str = ai_response[json_start:json_end]
                    optimized_data = json.loads(json_str)
                    
                    # Add metadata
                    optimized_data["source"] = "ai"
                    optimized_data["optimization_timestamp"] = datetime.now().isoformat()
                    
                    return optimized_data
                
            except json.JSONDecodeError as e:
                self.log_action("WARNING", f"Failed to parse AI response as JSON: {str(e)}")
                
            # Fallback to basic optimization
            return self._generate_basic_optimized_content(resume_analysis, skill_gaps)
            
        except Exception as e:
            self.log_action("ERROR", f"AI optimization failed: {str(e)}")
            return self._generate_basic_optimized_content(resume_analysis, skill_gaps)
    
    def _generate_basic_optimized_content(self, resume_analysis: Dict[str, Any], 
                                        skill_gaps: Dict[str, Any]) -> Dict[str, Any]:
        """Generate basic optimized content without AI."""
        
        current_skills = [skill["skill"] for skill in resume_analysis.get("current_skills", [])]
        missing_skills = [skill["skill"] for skill in skill_gaps.get("missing_skills", [])]
        
        # Create basic optimization
        optimized_summary = resume_analysis.get("summary", "")
        if missing_skills:
            optimized_summary += f" Skilled in {', '.join(missing_skills[:3])} and related technologies."
        
        return {
            "optimized_summary": optimized_summary,
            "enhanced_experience": resume_analysis.get("experience_details", []),
            "skills_section": {
                "technical_skills": current_skills + missing_skills[:5],
                "core_competencies": ["Problem Solving", "Team Collaboration", "Continuous Learning"]
            },
            "keywords_added": missing_skills[:5],
            "modifications_summary": [f"Added {len(missing_skills)} missing skills", "Enhanced summary"],
            "ats_optimization_tips": ["Include relevant keywords", "Use action verbs", "Quantify achievements"],
            "source": "basic",
            "optimization_timestamp": datetime.now().isoformat()
        }
    
    def _create_modification_plan(self, resume_analysis: Dict[str, Any], 
                                skill_gaps: Dict[str, Any],
                                optimized_content: Dict[str, Any],
                                job_title: str) -> Dict[str, Any]:
        """Create a detailed plan for resume modifications."""
        
        plan = {
            "summary": "Resume modification plan to better align with job requirements",
            "priority_changes": [],
            "skill_additions": [],
            "content_enhancements": [],
            "ats_optimizations": [],
            "estimated_time": "15-30 minutes"
        }
        
        # Priority changes based on skill gaps
        missing_skills = skill_gaps.get("missing_skills", [])
        high_priority_missing = [skill for skill in missing_skills if skill.get("priority") == "high"]
        
        if high_priority_missing:
            plan["priority_changes"].append({
                "type": "add_skills",
                "description": f"Add {len(high_priority_missing)} high-priority missing skills",
                "skills": [skill["skill"] for skill in high_priority_missing],
                "priority": "high"
            })
        
        # Content enhancements
        if optimized_content.get("optimized_summary"):
            plan["content_enhancements"].append({
                "type": "update_summary",
                "description": "Update professional summary to highlight relevant experience",
                "priority": "high"
            })
        
        # ATS optimizations
        plan["ats_optimizations"].extend([
            {
                "type": "keyword_optimization",
                "description": "Incorporate job-specific keywords naturally",
                "priority": "medium"
            },
            {
                "type": "format_optimization", 
                "description": "Ensure clean, scannable format for ATS systems",
                "priority": "medium"
            }
        ])
        
        return plan
    
    def _generate_ats_recommendations(self, resume_analysis: Dict[str, Any], 
                                    job_description: str, 
                                    skill_gaps: Dict[str, Any]) -> List[str]:
        """Generate ATS optimization recommendations."""
        
        recommendations = []
        
        # Keyword optimization
        missing_skills = skill_gaps.get("missing_skills", [])
        if missing_skills:
            top_missing = [skill["skill"] for skill in missing_skills[:3]]
            recommendations.append(f"Incorporate these keywords naturally: {', '.join(top_missing)}")
        
        # Format recommendations
        recommendations.extend([
            "Use standard section headings (Experience, Education, Skills)",
            "Avoid tables, graphics, or complex formatting",
            "Use standard fonts (Arial, Calibri, Times New Roman)",
            "Keep file size under 2MB",
            "Use bullet points for achievements and responsibilities"
        ])
        
        # Content recommendations
        recommendations.extend([
            "Quantify achievements with numbers and percentages",
            "Use action verbs at the beginning of bullet points",
            "Include relevant certifications and training",
            "Tailor experience descriptions to match job requirements"
        ])
        
        return recommendations
    
    async def create_modified_resume(self, modification_result: Dict[str, Any], 
                                   original_resume_path: str,
                                   job_title: str, 
                                   company_name: str) -> Optional[str]:
        """Create a modified resume file based on the modification results."""
        
        try:
            # Create output directory
            output_dir = os.path.join(os.path.dirname(original_resume_path), "modified_resumes")
            os.makedirs(output_dir, exist_ok=True)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_job_title = re.sub(r'[^\w\s-]', '', job_title).replace(' ', '_')[:30]
            safe_company = re.sub(r'[^\w\s-]', '', company_name).replace(' ', '_')[:20]
            
            filename = f"resume_{safe_job_title}_{safe_company}_{timestamp}.docx"
            output_path = os.path.join(output_dir, filename)
            
            # Create the modified document
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'Resume - {job_title}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add modification info
            doc.add_paragraph(f'Modified for: {job_title} at {company_name}')
            doc.add_paragraph(f'Modification date: {datetime.now().strftime("%B %d, %Y")}')
            doc.add_paragraph('')
            
            # Add optimized content
            optimized_content = modification_result.get("optimized_content", {})
            
            # Summary section
            if optimized_content.get("optimized_summary"):
                doc.add_heading('Professional Summary', level=1)
                doc.add_paragraph(optimized_content["optimized_summary"])
                doc.add_paragraph('')
            
            # Skills section
            skills_section = optimized_content.get("skills_section", {})
            if skills_section:
                doc.add_heading('Technical Skills', level=1)
                
                if skills_section.get("technical_skills"):
                    doc.add_paragraph(', '.join(skills_section["technical_skills"]))
                
                if skills_section.get("core_competencies"):
                    doc.add_paragraph('')
                    doc.add_paragraph('Core Competencies:')
                    doc.add_paragraph(', '.join(skills_section["core_competencies"]))
                
                doc.add_paragraph('')
            
            # Experience section (enhanced)
            enhanced_experience = optimized_content.get("enhanced_experience", [])
            if enhanced_experience:
                doc.add_heading('Professional Experience', level=1)
                
                for exp in enhanced_experience:
                    # Job title and company
                    job_header = doc.add_heading(f'{exp.get("title", "Job Title")} - {exp.get("company", "Company")}', level=2)
                    
                    # Description
                    if exp.get("description"):
                        doc.add_paragraph(exp["description"])
                    
                    # Achievements
                    if exp.get("achievements"):
                        for achievement in exp["achievements"]:
                            p = doc.add_paragraph()
                            p.add_run('• ').bold = True
                            p.add_run(achievement)
                    
                    doc.add_paragraph('')
            
            # Add modification summary
            modifications = modification_result.get("modification_plan", {}).get("priority_changes", [])
            if modifications:
                doc.add_heading('Modifications Made', level=1)
                for mod in modifications:
                    doc.add_paragraph(f'• {mod["description"]}')
                doc.add_paragraph('')
            
            # Add ATS recommendations
            ats_recs = modification_result.get("ats_recommendations", [])
            if ats_recs:
                doc.add_heading('ATS Optimization Tips', level=1)
                for rec in ats_recs[:5]:  # Top 5 recommendations
                    doc.add_paragraph(f'• {rec}')
            
            # Save the document
            doc.save(output_path)
            
            self.log_action("SUCCESS", f"Modified resume saved to: {output_path}")
            return output_path
            
        except Exception as e:
            self.log_action("ERROR", f"Failed to create modified resume: {str(e)}")
            return None
    
    async def get_resume_modification_tools(self) -> List[Dict[str, Any]]:
        """Get available tools for resume modification."""
        
        return [
            {
                "name": "modify_resume_for_job",
                "description": "Modify resume to match specific job requirements",
                "parameters": {
                    "resume_analysis": "Resume analysis data",
                    "job_title": "Target job title",
                    "job_description": "Job description text",
                    "company_name": "Company name"
                }
            },
            {
                "name": "analyze_skill_gaps",
                "description": "Analyze gaps between current skills and job requirements",
                "parameters": {
                    "resume_analysis": "Resume analysis data",
                    "job_description": "Job description text"
                }
            },
            {
                "name": "generate_ats_recommendations",
                "description": "Generate ATS optimization recommendations",
                "parameters": {
                    "resume_analysis": "Resume analysis data",
                    "job_description": "Job description text"
                }
            },
            {
                "name": "create_modified_resume",
                "description": "Create a modified resume file",
                "parameters": {
                    "modification_result": "Resume modification results",
                    "original_resume_path": "Path to original resume",
                    "job_title": "Target job title",
                    "company_name": "Company name"
                }
            }
        ]
