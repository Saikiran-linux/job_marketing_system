import asyncio
import re
from typing import Dict, Any, List, Set
from docx import Document
import os
from datetime import datetime
import openai
from openai import AsyncOpenAI
from agents.base_agent import BaseAgent, AgentState
from config import Config

class ResumeAnalysisAgent(BaseAgent):
    """Agent responsible for analyzing resumes and extracting current skills and content."""
    
    def __init__(self):
        super().__init__("ResumeAnalysisAgent")
        self.client = AsyncOpenAI(api_key=Config.OPENAI_API_KEY) if Config.OPENAI_API_KEY else None
        
        # Common sections in resumes
        self.resume_sections = {
            "contact": ["contact", "personal", "information"],
            "summary": ["summary", "objective", "profile", "about"],
            "experience": ["experience", "work", "employment", "professional"],
            "education": ["education", "academic", "qualification"],
            "skills": ["skills", "technical", "competencies", "technologies"],
            "projects": ["projects", "portfolio", "accomplishments"],
            "certifications": ["certifications", "certificates", "licenses"],
            "awards": ["awards", "honors", "achievements", "recognition"]
        }
    
    async def execute(self, state: AgentState) -> AgentState:
        """Analyze resume and extract current skills and content structure."""
        
        # Validate required inputs
        required_fields = ["resume_path"]
        if not self.validate_input(state, required_fields):
            state.status = "error"
            state.error = "Missing resume path"
            return state
        
        resume_path = state.resume_path
        
        if not os.path.exists(resume_path):
            state.status = "error"
            state.error = f"Resume file not found: {resume_path}"
            return state
        
        self.log_action("ANALYZING", f"Resume: {os.path.basename(resume_path)}")
        
        try:
            # Extract text content from resume
            resume_content = await self._extract_resume_content(resume_path)
            
            if not resume_content:
                state.status = "error"
                state.error = "Could not extract content from resume"
                return state
            
            # Analyze resume structure
            sections = self._identify_sections(resume_content)
            
            # Extract current skills
            current_skills = self._extract_current_skills(resume_content)
            
            # Extract experience details
            experience_details = self._extract_experience_details(resume_content)
            
            # Extract education details
            education_details = self._extract_education_details(resume_content)
            
            # Extract contact information
            contact_info = self._extract_contact_info(resume_content)
            
            # Analyze resume quality using AI (if available)
            quality_analysis = await self._analyze_resume_quality(resume_content)
            
            # Calculate resume statistics
            stats = self._calculate_resume_stats(resume_content, sections)
            
            # Update state with results
            state.resume_analysis = {
                "status": "success",
                "resume_content": resume_content,
                "sections": sections,
                "current_skills": current_skills,
                "experience_details": experience_details,
                "education_details": education_details,
                "contact_info": contact_info,
                "quality_analysis": quality_analysis,
                "statistics": stats,
                "file_path": resume_path,
                "analysis_timestamp": datetime.now().isoformat()
            }
            
            state.steps_completed.append("resume_analysis")
            state.current_step = "resume_analysis_complete"
            
            self.log_action("SUCCESS", f"Extracted {len(current_skills)} skills from resume")
            return state
            
        except Exception as e:
            self.log_action("ERROR", f"Resume analysis failed: {str(e)}")
            state.status = "error"
            state.error = f"Resume analysis failed: {str(e)}"
            return state
    
    async def _extract_resume_content(self, resume_path: str) -> str:
        """Extract text content from resume file."""
        
        file_extension = os.path.splitext(resume_path)[1].lower()
        
        try:
            if file_extension == '.docx':
                doc = Document(resume_path)
                content = []
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            content.append(" | ".join(row_text))
                
                return "\n".join(content)
                
            elif file_extension == '.txt':
                with open(resume_path, 'r', encoding='utf-8') as file:
                    return file.read()
                    
            elif file_extension == '.pdf':
                # For PDF files, you'd need a PDF extraction library like PyPDF2 or pdfplumber
                self.log_action("WARNING", "PDF extraction requires additional dependencies")
                return ""
                
            else:
                self.log_action("ERROR", f"Unsupported file format: {file_extension}")
                return ""
                
        except Exception as e:
            self.log_action("ERROR", f"Failed to extract content: {str(e)}")
            return ""
    
    def _identify_sections(self, resume_content: str) -> Dict[str, Dict[str, Any]]:
        """Identify different sections in the resume."""
        
        sections = {}
        lines = resume_content.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line_stripped = line.strip().lower()
            
            if not line_stripped:
                continue
            
            # Check if this line is a section header
            for section_name, keywords in self.resume_sections.items():
                if any(keyword in line_stripped for keyword in keywords):
                    # Check if this looks like a header (short line, possibly uppercase)
                    if len(line.strip()) < 50 and (line.isupper() or line.istitle()):
                        current_section = section_name
                        sections[section_name] = {
                            "start_line": i,
                            "header": line.strip(),
                            "content": []
                        }
                        break
            
            # Add content to current section
            if current_section and current_section in sections:
                sections[current_section]["content"].append(line.strip())
        
        # Set end lines for sections
        section_names = list(sections.keys())
        for i, section_name in enumerate(section_names):
            if i < len(section_names) - 1:
                next_section = section_names[i + 1]
                sections[section_name]["end_line"] = sections[next_section]["start_line"] - 1
            else:
                sections[section_name]["end_line"] = len(lines) - 1
        
        return sections
    
    def _extract_current_skills(self, resume_content: str) -> List[Dict[str, Any]]:
        """Extract skills mentioned in the resume."""
        
        # Technical skills patterns
        skill_patterns = {
            "programming_languages": [
                "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
                "kotlin", "swift", "php", "ruby", "scala", "r", "matlab", "sql"
            ],
            "frameworks": [
                "react", "angular", "vue", "django", "flask", "fastapi", "spring", "express",
                "node.js", "next.js", "nuxt.js", "laravel", "rails", "asp.net", "tensorflow",
                "pytorch", "scikit-learn", "pandas", "numpy"
            ],
            "databases": [
                "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "cassandra",
                "oracle", "sqlite", "dynamodb", "neo4j"
            ],
            "cloud_platforms": [
                "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "terraform",
                "ansible", "jenkins", "gitlab", "github actions"
            ],
            "tools": [
                "git", "jira", "confluence", "slack", "teams", "figma", "sketch", "photoshop",
                "tableau", "power bi", "excel", "powerpoint"
            ]
        }
        
        found_skills = []
        resume_lower = resume_content.lower()
        
        for category, skills in skill_patterns.items():
            for skill in skills:
                # Count mentions
                mention_count = resume_lower.count(skill.lower())
                
                if mention_count > 0:
                    # Find context where skill is mentioned
                    contexts = []
                    for match in re.finditer(re.escape(skill.lower()), resume_lower):
                        start = max(0, match.start() - 50)
                        end = min(len(resume_lower), match.end() + 50)
                        context = resume_content[start:end].strip()
                        contexts.append(context)
                    
                    found_skills.append({
                        "skill": skill,
                        "category": category,
                        "mention_count": mention_count,
                        "contexts": contexts[:3],  # Keep only first 3 contexts
                        "confidence": min(1.0, mention_count * 0.3)
                    })
        
        # Sort by mention count and confidence
        found_skills.sort(key=lambda x: (x["mention_count"], x["confidence"]), reverse=True)
        
        return found_skills
    
    def _extract_experience_details(self, resume_content: str) -> List[Dict[str, Any]]:
        """Extract work experience details."""
        
        experience_entries = []
        
        # Common patterns for experience entries
        date_patterns = [
            r'(\d{4})\s*[-–—]\s*(\d{4}|\w+)',  # 2020 - 2023 or 2020 - Present
            r'(\w+\s+\d{4})\s*[-–—]\s*(\w+\s+\d{4}|\w+)',  # Jan 2020 - Dec 2023
        ]
        
        # Split content into potential experience entries
        lines = resume_content.split('\n')
        current_entry = None
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                continue
            
            # Check for date patterns (likely indicates start of experience entry)
            for pattern in date_patterns:
                date_match = re.search(pattern, line_stripped)
                if date_match:
                    # Save previous entry if exists
                    if current_entry:
                        experience_entries.append(current_entry)
                    
                    # Start new entry
                    current_entry = {
                        "date_range": date_match.group(0),
                        "start_date": date_match.group(1),
                        "end_date": date_match.group(2),
                        "full_text": [line_stripped],
                        "title": "",
                        "company": "",
                        "description": []
                    }
                    break
            
            # Add to current entry
            if current_entry:
                current_entry["full_text"].append(line_stripped)
                
                # Try to identify title and company
                if not current_entry["title"] and len(current_entry["full_text"]) <= 3:
                    # Assume title is in first few lines after date
                    if "at" in line_stripped.lower() or "|" in line_stripped:
                        parts = re.split(r'\s+at\s+|\s*\|\s*', line_stripped, 1)
                        if len(parts) == 2:
                            current_entry["title"] = parts[0].strip()
                            current_entry["company"] = parts[1].strip()
                    elif not current_entry["title"]:
                        current_entry["title"] = line_stripped
                
                # Add to description if it's detailed content
                elif len(line_stripped) > 20:
                    current_entry["description"].append(line_stripped)
        
        # Add last entry
        if current_entry:
            experience_entries.append(current_entry)
        
        return experience_entries
    
    def _extract_education_details(self, resume_content: str) -> List[Dict[str, Any]]:
        """Extract education details."""
        
        education_entries = []
        
        # Education degree patterns
        degree_patterns = [
            r'\b(?:bachelor\'?s?|ba|bs|b\.a\.|b\.s\.)\s+(?:of\s+|in\s+)?([^\n,]+)',
            r'\b(?:master\'?s?|ma|ms|m\.a\.|m\.s\.|mba)\s+(?:of\s+|in\s+)?([^\n,]+)',
            r'\b(?:phd|ph\.d\.|doctorate|doctoral)\s+(?:of\s+|in\s+)?([^\n,]+)',
        ]
        
        for pattern in degree_patterns:
            matches = re.finditer(pattern, resume_content, re.IGNORECASE)
            for match in matches:
                # Extract surrounding context
                start = max(0, match.start() - 100)
                end = min(len(resume_content), match.end() + 100)
                context = resume_content[start:end]
                
                education_entries.append({
                    "degree": match.group(0),
                    "field": match.group(1) if len(match.groups()) > 0 else "",
                    "context": context.strip()
                })
        
        return education_entries
    
    def _extract_contact_info(self, resume_content: str) -> Dict[str, str]:
        """Extract contact information from resume."""
        
        contact_info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, resume_content)
        if email_match:
            contact_info["email"] = email_match.group(0)
        
        # Phone pattern
        phone_patterns = [
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, resume_content)
            if phone_match:
                contact_info["phone"] = phone_match.group(0)
                break
        
        # LinkedIn profile
        linkedin_pattern = r'linkedin\.com/in/[A-Za-z0-9-]+'
        linkedin_match = re.search(linkedin_pattern, resume_content, re.IGNORECASE)
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group(0)
        
        # GitHub profile
        github_pattern = r'github\.com/[A-Za-z0-9-]+'
        github_match = re.search(github_pattern, resume_content, re.IGNORECASE)
        if github_match:
            contact_info["github"] = github_match.group(0)
        
        return contact_info
    
    async def _analyze_resume_quality(self, resume_content: str) -> Dict[str, Any]:
        """Analyze resume quality using AI."""
        
        if not self.client:
            return {"status": "skipped", "reason": "OpenAI API not configured"}
        
        try:
            prompt = f"""
            Analyze this resume and provide feedback on its quality and areas for improvement.
            
            Resume Content:
            {resume_content[:3000]}  # Limit content to avoid token limits
            
            Please provide a JSON response with the following structure:
            {{
                "overall_score": 7.5,
                "strengths": ["strength1", "strength2", ...],
                "weaknesses": ["weakness1", "weakness2", ...],
                "suggestions": ["suggestion1", "suggestion2", ...],
                "missing_sections": ["section1", "section2", ...],
                "formatting_issues": ["issue1", "issue2", ...],
                "keyword_optimization": {{
                    "score": 6.0,
                    "missing_keywords": ["keyword1", "keyword2", ...]
                }}
            }}
            
            Focus on:
            - Content quality and relevance
            - Structure and organization
            - Missing information
            - ATS compatibility
            - Professional presentation
            """
            
            response = await self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an expert resume reviewer and career coach. Provide constructive, actionable feedback."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=800,
                temperature=0.3
            )
            
            # Parse the AI response
            ai_response = response.choices[0].message.content
            
            # Extract JSON from the response
            import json
            try:
                json_start = ai_response.find('{')
                json_end = ai_response.rfind('}') + 1
                
                if json_start != -1 and json_end != -1:
                    json_str = ai_response[json_start:json_end]
                    quality_data = json.loads(json_str)
                    quality_data["status"] = "success"
                    return quality_data
                
            except json.JSONDecodeError:
                pass
            
            # If JSON parsing fails, return basic analysis
            return {
                "status": "partial",
                "ai_feedback": ai_response,
                "overall_score": 5.0
            }
            
        except Exception as e:
            self.log_action("ERROR", f"AI quality analysis failed: {str(e)}")
            return {
                "status": "error",
                "message": str(e)
            }
    
    def _calculate_resume_stats(self, resume_content: str, sections: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate resume statistics."""
        
        lines = resume_content.split('\n')
        words = resume_content.split()
        
        return {
            "total_lines": len(lines),
            "total_words": len(words),
            "total_characters": len(resume_content),
            "sections_found": len(sections),
            "section_names": list(sections.keys()),
            "average_words_per_line": len(words) / len(lines) if lines else 0,
            "estimated_pages": max(1, len(lines) // 40)  # Rough estimate
        }
